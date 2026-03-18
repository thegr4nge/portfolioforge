"""Word document export for analysis reports.

Produces a professional .docx file suitable for attaching to SMSF client files.
The document renders the same data as the terminal report in a format that
accountants can annotate, brand, and archive.

Usage::

    from market_data.analysis.exporter import export_report
    export_report(report, conn, Path("client_smith.docx"))
"""

from __future__ import annotations

import sqlite3
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from market_data.analysis.models import AnalysisReport
from market_data.analysis.narrative import (
    DISCLAIMER,
    narrative_cagr,
    narrative_max_drawdown,
    narrative_sharpe,
    narrative_total_return,
)
from market_data.analysis.scenario import compute_drawdown_series, compute_recovery_days
from market_data.backtest.models import BacktestResult
from market_data.backtest.tax.audit import build_cgt_event_rows
from market_data.backtest.tax.audit_models import CgtEventRow
from market_data.backtest.tax.engine import TAX_ENGINE_VERSION
from market_data.backtest.tax.franking import resolve_franking_pct
from market_data.backtest.tax.models import DisposedLot, TaxAwareResult, TaxSummary
from market_data.backtest.tax.trade_record import TradeRecord
from market_data.verification.workpaper_id import generate_workpaper_id

# ── colour palette ──────────────────────────────────────────────────────────────
_NAVY_HEX = "1F3864"
_TEAL_HEX = "00878A"
_GRAY_HEX = "F2F2F2"

_NAVY = RGBColor(0x1F, 0x38, 0x64)
_TEAL = RGBColor(0x00, 0x87, 0x8A)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_RED = RGBColor(0xC0, 0x00, 0x00)


# ── low-level helpers ───────────────────────────────────────────────────────────


def _shade_cell(cell: object, hex_color: str) -> None:
    """Set table cell background fill colour (hex string without #)."""
    tc = cell._tc  # type: ignore[attr-defined]
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _header_row(table: object, labels: list[str], *, font_size: int = 9) -> None:
    """Write a navy-shaded header row with white bold text."""
    row = table.rows[0]  # type: ignore[attr-defined]
    for cell, label in zip(row.cells, labels, strict=False):
        _shade_cell(cell, _NAVY_HEX)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(label)
        run.bold = True
        run.font.color.rgb = _WHITE
        run.font.size = Pt(font_size)


def _body_row(
    table: object,
    values: list[str],
    *,
    shade: bool = False,
    font_size: int = 9,
) -> None:
    """Append a data row; optionally shades in light grey."""
    row = table.add_row()  # type: ignore[attr-defined]
    for cell, value in zip(row.cells, values, strict=False):
        if shade:
            _shade_cell(cell, _GRAY_HEX)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(value)
        run.font.size = Pt(font_size)


def _set_col_widths(table: object, widths_cm: list[float]) -> None:
    """Apply fixed column widths (cm) to every cell in a table.

    Prevents Word from auto-fitting content, which causes overflow on
    wide tables. Total widths_cm should sum to the usable page width.
    """
    for row in table.rows:  # type: ignore[attr-defined]
        for cell, w in zip(row.cells, widths_cm, strict=False):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcW")):
                tcPr.remove(existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(w * 567)))  # cm → twips (1 cm ≈ 567 twips)
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _add_cell_borders(table: object) -> None:
    """Add thin light-grey borders to every cell in a table."""
    for row in table.rows:  # type: ignore[attr-defined]
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "BFBFBF")
                borders.append(el)
            tcPr.append(borders)


def _section_heading(doc: object, text: str) -> None:
    """Add a teal bold section heading followed by a thin rule."""
    para = doc.add_paragraph()  # type: ignore[attr-defined]
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = _TEAL
    run.font.size = Pt(13)


def _get_backtest(report: AnalysisReport) -> BacktestResult:
    result = report.result
    if isinstance(result, TaxAwareResult):
        return result.backtest
    return result


# ── audit annotations (PROF-01 / PROF-02) ───────────────────────────────────────

# Static ATO elections displayed in the Calculation Methodology table (PROF-02).
# Tuple: (Rule, Election / Method, ATO Reference)
_CGT_RULES: list[tuple[str, str, str]] = [
    ("Cost basis method", "First In, First Out (FIFO)", "ITAA 1997 s. 104-240"),
    ("CGT discount", "50% for assets held strictly > 12 months (individual)", "ITAA 1997 s. 115-100"),
    (
        "Discount threshold",
        "Disposed strictly after 12-month anniversary of acquisition",
        "ITAA 1997 s. 115-25",
    ),
    (
        "Loss ordering",
        "Losses applied to non-discountable gains first, then discountable gains "
        "before 50% discount",
        "ATO CGT guide",
    ),
    (
        "Franking credits",
        "45-day holding rule per dividend; $5,000 annual threshold waives rule",
        "ITAA 1936 s. 160APHO",
    ),
    (
        "Capital loss carry-forward",
        "Net losses carry forward indefinitely until absorbed",
        "ITAA 1997 s. 102-10",
    ),
    ("Australian tax year", "1 July to 30 June", "ITAA 1997 s. 995-1"),
    (
        "Brokerage treatment",
        "Acquisition brokerage added to cost base; disposal brokerage deducted from proceeds",
        "ITAA 1997 s. 110-25",
    ),
]


def _cgt_rule_annotation(row: CgtEventRow, entity_type: str = "individual") -> str:
    """One-line ATO rule annotation for a single CGT event (PROF-01)."""
    acq = row.acquired_date.strftime("%d %b %Y")
    months = (row.disposed_date.year - row.acquired_date.year) * 12 + (
        row.disposed_date.month - row.acquired_date.month
    )
    if row.gain_type == "discountable_gain":
        discount_label = "33.33% CGT discount applied (SMSF)" if entity_type == "smsf" else "50% CGT discount applied"
        return f"FIFO parcel {acq} — {discount_label} (held {months} months)"
    if row.gain_type == "non_discountable_gain":
        return f"FIFO parcel {acq} — no discount (held {months} months, \u226412)"
    return f"FIFO parcel {acq} — capital loss (held {months} months)"


# ── section builders ────────────────────────────────────────────────────────────


def _add_cover(
    doc: object,
    br: BacktestResult,
    after_tax_cagr: float | None = None,
    *,
    sample_data: bool = False,
    workpaper_id: str = "",
) -> None:
    """Cover page: title, period, 4 KPI boxes, optional after-tax CAGR line.

    Args:
        doc: python-docx Document object.
        br: BacktestResult supplying portfolio, metrics, and curve dates.
        after_tax_cagr: After-tax CAGR to display instead of pre-tax CAGR, or None.
        sample_data: When True, adds a visible "SAMPLE DATA" banner below the title.
    """
    title_para = doc.add_paragraph()  # type: ignore[attr-defined]
    title_run = title_para.add_run("PortfolioForge — Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = _NAVY

    if sample_data:
        banner_para = doc.add_paragraph()  # type: ignore[attr-defined]
        banner_run = banner_para.add_run("*** SAMPLE DATA — NOT BASED ON A REAL CLIENT RUN ***")
        banner_run.bold = True
        banner_run.font.size = Pt(12)
        banner_run.font.color.rgb = _RED

    tickers = ", ".join(f"{t} ({w:.0%})" for t, w in br.portfolio.items())
    sub = doc.add_paragraph()  # type: ignore[attr-defined]
    sub.add_run(f"Portfolio: {tickers}").font.size = Pt(10)

    period = doc.add_paragraph()  # type: ignore[attr-defined]
    start = br.equity_curve.index[0]
    end = br.equity_curve.index[-1]
    period.add_run(
        f"Period: {start.strftime('%d %b %Y')} – {end.strftime('%d %b %Y')}"
        f"   |   Generated: {date.today().strftime('%d %b %Y')}"
    ).font.size = Pt(10)

    if workpaper_id:
        vid_para = doc.add_paragraph()  # type: ignore[attr-defined]
        vid_run = vid_para.add_run(f"Verification ID: {workpaper_id}")
        vid_run.font.size = Pt(9)
        vid_run.font.color.rgb = _NAVY

    doc.add_paragraph()  # type: ignore[attr-defined]

    # 4 KPI boxes as a 2×4 table (label / value / label / value)
    # When tax data is present, replace "CAGR" with "After-Tax CAGR" — the
    # headline number for SMSF accountants. Pre-tax CAGR remains in the
    # Performance vs Benchmark table below.
    cagr_label = "After-Tax CAGR" if after_tax_cagr is not None else "CAGR"
    cagr_value = f"{after_tax_cagr:.1%}" if after_tax_cagr is not None else f"{br.metrics.cagr:.1%}"
    kpi_table = doc.add_table(rows=2, cols=4)  # type: ignore[attr-defined]
    kpi_labels = ["Total Return", cagr_label, "Max Drawdown", "Sharpe Ratio"]
    kpi_values = [
        f"{br.metrics.total_return:.1%}",
        cagr_value,
        f"{br.metrics.max_drawdown:.1%}",
        f"{br.metrics.sharpe_ratio:.2f}",
    ]
    label_row = kpi_table.rows[0]
    value_row = kpi_table.rows[1]
    for i, (lbl, val) in enumerate(zip(kpi_labels, kpi_values, strict=True)):
        _shade_cell(label_row.cells[i], _NAVY_HEX)
        lp = label_row.cells[i].paragraphs[0]
        lp.clear()
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.color.rgb = _WHITE
        lr.font.size = Pt(9)

        _shade_cell(value_row.cells[i], _TEAL_HEX)
        vp = value_row.cells[i].paragraphs[0]
        vp.clear()
        vr = vp.add_run(val)
        vr.bold = True
        vr.font.color.rgb = _WHITE
        vr.font.size = Pt(14)

    doc.add_page_break()  # type: ignore[attr-defined]


def _add_composition(doc: object, br: BacktestResult) -> None:
    """Portfolio composition table: ticker, weight, franking %."""
    _section_heading(doc, "Portfolio Composition")

    table = doc.add_table(rows=1, cols=3)  # type: ignore[attr-defined]
    _header_row(table, ["Ticker", "Weight", "Franking Credit %"])

    for i, (ticker, weight) in enumerate(br.portfolio.items()):
        franking_pct = resolve_franking_pct(ticker, None)
        _body_row(
            table,
            [ticker, f"{weight:.1%}", f"{franking_pct:.0%}"],
            shade=(i % 2 == 1),
        )
    _add_cell_borders(table)
    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_performance(doc: object, br: BacktestResult) -> None:
    """Performance vs benchmark table + narrative."""
    _section_heading(doc, "Performance vs Benchmark")

    table = doc.add_table(rows=1, cols=3)  # type: ignore[attr-defined]
    _header_row(table, ["Metric", "Portfolio", br.benchmark.ticker])

    rows = [
        ("Total Return", f"{br.metrics.total_return:.2%}", f"{br.benchmark.total_return:.2%}"),
        ("CAGR", f"{br.metrics.cagr:.2%}", f"{br.benchmark.cagr:.2%}"),
        ("Max Drawdown", f"{br.metrics.max_drawdown:.2%}", f"{br.benchmark.max_drawdown:.2%}"),
        ("Sharpe Ratio", f"{br.metrics.sharpe_ratio:.2f}", f"{br.benchmark.sharpe_ratio:.2f}"),
    ]
    for i, (metric, port_val, bench_val) in enumerate(rows):
        _body_row(table, [metric, port_val, bench_val], shade=(i % 2 == 1))
    _add_cell_borders(table)
    doc.add_paragraph()  # type: ignore[attr-defined]

    dd_series = compute_drawdown_series(br.equity_curve)
    max_dd = float(dd_series.min()) * 100
    recovery = compute_recovery_days(br.equity_curve)

    narratives = [
        narrative_total_return(br.metrics.total_return * 100),
        narrative_cagr(br.metrics.cagr * 100),
        narrative_max_drawdown(max_dd, recovery),
        narrative_sharpe(br.metrics.sharpe_ratio),
    ]
    for sentence in narratives:
        p = doc.add_paragraph(style="List Bullet")  # type: ignore[attr-defined]
        p.add_run(sentence).font.size = Pt(9)

    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_tax_analysis(doc: object, tax_result: TaxAwareResult) -> None:
    """Year-by-year Australian tax summary table."""
    _section_heading(doc, "Australian Tax Analysis (CGT)")

    tax = tax_result.tax
    table = doc.add_table(rows=1, cols=6)  # type: ignore[attr-defined]
    _header_row(
        table,
        [
            "Tax Year",
            "CGT Events",
            "CGT Payable (AUD)",
            "Franking Credits (AUD)",
            "Dividend Income (AUD)",
            "Carry-Fwd Loss (AUD)",
        ],
    )

    for i, yr in enumerate(tax.years):
        _body_row(
            table,
            [
                f"FY{yr.ending_year}",
                str(yr.cgt_events),
                f"${yr.cgt_payable:,.2f}",
                f"${yr.franking_credits_claimed:,.2f}",
                f"${yr.dividend_income:,.2f}",
                f"${yr.carried_forward_loss:,.2f}",
            ],
            shade=(i % 2 == 1),
        )
    _add_cell_borders(table)

    p = doc.add_paragraph()  # type: ignore[attr-defined]
    p.add_run(
        f"Total tax paid: ${tax.total_tax_paid:,.2f}   |   "
        f"After-tax CAGR: {tax.after_tax_cagr:.2%}"
    ).font.size = Pt(9)

    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_cgt_log(doc: object, tax_result: TaxAwareResult, entity_type: str = "individual") -> None:
    """Individual CGT event log using structured audit rows.

    7 columns at fixed widths to fit A4 portrait (16 cm usable):
    #(0.4) | Tax Year(1.5) | Ticker(1.5) | Acquired(2.0) | Disposed(2.0)
    | Gain/Loss(2.1) | ATO Rule Applied(6.5)
    Cost base and proceeds are omitted here; they appear in the full audit
    export. Gain = proceeds − cost base for any manual cross-check.
    """
    _section_heading(doc, "CGT Event Log")

    rows = build_cgt_event_rows(tax_result.tax.lots)
    if not rows:
        doc.add_paragraph("No disposal events recorded in this period.")  # type: ignore[attr-defined]
        return

    table = doc.add_table(rows=1, cols=7)  # type: ignore[attr-defined]
    _header_row(
        table,
        [
            "#",
            "Tax Year",
            "Ticker",
            "Acquired",
            "Disposed",
            "Gain/Loss (AUD)",
            "ATO Rule Applied",
        ],
        font_size=8,
    )

    for i, row in enumerate(rows):
        _body_row(
            table,
            [
                str(i + 1),
                row.tax_year_label,
                row.ticker,
                row.acquired_date.strftime("%d %b %Y"),
                row.disposed_date.strftime("%d %b %Y"),
                f"${row.gain_aud:,.2f}",
                _cgt_rule_annotation(row, entity_type),
            ],
            shade=(i % 2 == 1),
            font_size=8,
        )

    _set_col_widths(table, [0.4, 1.5, 1.5, 2.0, 2.0, 2.1, 6.5])
    _add_cell_borders(table)
    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_coverage(doc: object, br: BacktestResult) -> None:
    """Data coverage table: ticker, date range, record count, quality status.

    Quality status is always "validated" because the backtest engine applies a
    quality_flags=0 filter on all price queries — only records that passed the
    full ValidationSuite are used in the simulation.
    """
    _section_heading(doc, "Data Coverage")

    table = doc.add_table(rows=1, cols=5)  # type: ignore[attr-defined]
    _header_row(table, ["Ticker", "From", "To", "Records", "Quality Status"])

    for i, cov in enumerate(br.coverage):
        _body_row(
            table,
            [
                cov.ticker,
                cov.from_date.strftime("%d %b %Y"),
                cov.to_date.strftime("%d %b %Y"),
                str(cov.records),
                "validated",
            ],
            shade=(i % 2 == 1),
        )
    _add_cell_borders(table)

    note = doc.add_paragraph()  # type: ignore[attr-defined]
    note.add_run(
        "Quality status 'validated': price records passed all quality checks "
        "(quality_flags=0 filter applied). Only validated records are used in backtests."
    ).font.size = Pt(8)

    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_methodology(doc: object, tax_result: TaxAwareResult | None = None) -> None:
    """ATO elections table (PROF-02) and plain-language methodology narrative."""
    _section_heading(doc, "Calculation Methodology")

    entity_type = tax_result.tax.entity_type if tax_result is not None else "individual"
    is_smsf = entity_type == "smsf"

    # ATO elections & rules table — override CGT discount row for SMSF entity type
    rules = list(_CGT_RULES)
    for idx, (rule, detail, ref) in enumerate(rules):
        if rule == "CGT discount" and is_smsf:
            rules[idx] = (
                "CGT discount",
                "33.33% for assets held strictly > 12 months (SMSF — ATO s.115-100)",
                "ITAA 1997 s. 115-100",
            )
        if rule == "Franking credits" and is_smsf:
            rules[idx] = (
                "Franking credits",
                "45-day holding rule per dividend; $5,000 small-shareholder exemption does NOT apply to SMSFs",
                "ITAA 1936 s. 160APHO",
            )

    table = doc.add_table(rows=1, cols=3)  # type: ignore[attr-defined]
    _header_row(table, ["Rule", "Election / Method", "ATO Reference"])

    for i, (rule, detail, ref) in enumerate(rules):
        _body_row(table, [rule, detail, ref], shade=(i % 2 == 1))
    _add_cell_borders(table)

    # Dynamic rows drawn from the actual backtest when tax data is present.
    if tax_result is not None:
        n = len(_CGT_RULES)
        tax = tax_result.tax
        final_carry = tax.years[-1].carried_forward_loss if tax.years else 0.0
        _body_row(
            table,
            [
                "Marginal tax rate applied",
                f"{tax.marginal_tax_rate:.1%}",
                "User-supplied parameter",
            ],
            shade=(n % 2 == 1),
        )
        _body_row(
            table,
            [
                "Carry-forward loss at period end",
                f"${final_carry:,.2f}",
                "Carries to next tax year",
            ],
            shade=((n + 1) % 2 == 1),
        )
        _body_row(
            table,
            [
                "Tax engine version",
                TAX_ENGINE_VERSION,
                "PortfolioForge internal — audit traceability",
            ],
            shade=((n + 2) % 2 == 1),
        )

    doc.add_paragraph()  # type: ignore[attr-defined]

    text = (
        "This report is generated by PortfolioForge. Backtesting simulates historical portfolio "
        "performance using end-of-day price data with mandatory brokerage costs ($10 minimum or "
        "0.1% of trade value). Prices are adjusted for corporate actions (splits). "
        "CGT calculations are ATO-validated against published worked examples "
        "(Sonya short-term, Mei-Ling long-term, FIFO multi-parcel). "
        "This tool is not licensed to provide personal financial advice."
    )
    p = doc.add_paragraph()  # type: ignore[attr-defined]
    p.add_run(text).font.size = Pt(9)
    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_page_numbers(doc: object) -> None:
    """Add 'Page X of Y' footer to all sections via OOXML field codes."""
    for section in doc.sections:  # type: ignore[attr-defined]
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT = 2

        run = para.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = _NAVY

        # PAGE field
        fld_page = OxmlElement("w:fldChar")
        fld_page.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_page)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

        run2 = para.add_run(" of ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = _NAVY

        # NUMPAGES field
        fld_np = OxmlElement("w:fldChar")
        fld_np.set(qn("w:fldCharType"), "begin")
        run2._r.append(fld_np)
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = " NUMPAGES "
        run2._r.append(instr2)
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        run2._r.append(fld_end2)

        run3 = para.add_run("   |   PortfolioForge")
        run3.font.size = Pt(8)
        run3.font.color.rgb = _NAVY


def _add_disclaimer(doc: object) -> None:
    """Mandatory disclaimer — always the final section."""
    _section_heading(doc, "Disclaimer")
    p = doc.add_paragraph()  # type: ignore[attr-defined]
    run = p.add_run(DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _NAVY


# ── public entry point ──────────────────────────────────────────────────────────


def export_report(
    report: AnalysisReport,
    conn: sqlite3.Connection,
    output_path: Path,
    *,
    sample_data: bool = False,
) -> str:
    """Export an AnalysisReport to a .docx file.

    All sections matching the data available are written. Tax sections
    (Australian Tax Analysis, CGT Event Log) are included only when the
    report contains a TaxAwareResult; pre-tax-only reports skip them with
    a note.

    The mandatory DISCLAIMER is always the final section regardless of data.

    Args:
        report: AnalysisReport wrapping a BacktestResult or TaxAwareResult.
        conn: SQLite connection (used for sector/geo — reserved for future use).
        output_path: Destination .docx path. Parent directory must exist.
        sample_data: When True, adds a visible "SAMPLE DATA" banner on the cover
            page. Use this for demos and presentations not based on a real client
            run, to comply with the AFSL disclaimer requirement.

    Raises:
        ValueError: If output_path does not have a .docx extension.
        OSError: If output_path cannot be written.
    """
    if output_path.suffix.lower() != ".docx":
        raise ValueError(f"output_path must end in .docx, got: {output_path}")

    workpaper_id = generate_workpaper_id()
    br = _get_backtest(report)
    doc = Document()

    # Narrow default margins (2 cm) for a more professional look.
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    after_tax_cagr = (
        report.result.tax.after_tax_cagr if isinstance(report.result, TaxAwareResult) else None
    )
    _add_page_numbers(doc)
    _add_cover(doc, br, after_tax_cagr=after_tax_cagr, sample_data=sample_data, workpaper_id=workpaper_id)
    _add_composition(doc, br)
    _add_performance(doc, br)

    if isinstance(report.result, TaxAwareResult):
        _entity_type = report.result.tax.entity_type
        _add_tax_analysis(doc, report.result)
        _add_cgt_log(doc, report.result, entity_type=_entity_type)
    else:
        p = doc.add_paragraph()
        p.add_run(
            "Tax analysis not available — run with tax-aware backtest to include CGT summary."
        ).italic = True

    _add_coverage(doc, br)
    _add_methodology(
        doc,
        tax_result=report.result if isinstance(report.result, TaxAwareResult) else None,
    )
    _add_disclaimer(doc)

    doc.save(str(output_path))
    return workpaper_id


# ── CSV import workpaper helpers ────────────────────────────────────────────────


def _add_trade_history_table(doc: object, trades: list[TradeRecord]) -> None:
    """Actual broker trade history — 7 columns."""
    _section_heading(doc, "Trade History (Imported from Broker CSV)")

    table = doc.add_table(rows=1, cols=7)  # type: ignore[attr-defined]
    _header_row(
        table,
        ["Date", "Ticker", "Action", "Qty", "Price (AUD)", "Brokerage (AUD)", "Total (AUD)"],
        font_size=8,
    )
    for i, r in enumerate(sorted(trades, key=lambda t: t.trade_date)):
        total = r.quantity * r.price_aud + (r.brokerage_aud if r.action == "BUY" else -r.brokerage_aud)
        _body_row(
            table,
            [
                str(r.trade_date),
                r.ticker,
                r.action,
                f"{r.quantity:,.4g}",
                f"${r.price_aud:,.4f}",
                f"${r.brokerage_aud:,.2f}",
                f"${total:,.2f}",
            ],
            shade=(i % 2 == 1),
            font_size=8,
        )
    _set_col_widths(table, [2.0, 1.5, 1.2, 1.2, 2.0, 2.2, 2.0])
    _add_cell_borders(table)
    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_cgt_tax_summary(doc: object, tax: TaxSummary) -> None:
    """Year-by-year CGT summary table from a TaxSummary (no BacktestResult needed)."""
    _section_heading(doc, "Australian Tax Analysis (CGT)")

    table = doc.add_table(rows=1, cols=6)  # type: ignore[attr-defined]
    _header_row(
        table,
        [
            "Tax Year",
            "CGT Events",
            "CGT Payable (AUD)",
            "Franking Credits (AUD)",
            "Dividend Income (AUD)",
            "Carry-Fwd Loss (AUD)",
        ],
    )
    for i, yr in enumerate(tax.years):
        _body_row(
            table,
            [
                f"FY{yr.ending_year}",
                str(yr.cgt_events),
                f"${yr.cgt_payable:,.2f}",
                f"${yr.franking_credits_claimed:,.2f}",
                f"${yr.dividend_income:,.2f}",
                f"${yr.carried_forward_loss:,.2f}",
            ],
            shade=(i % 2 == 1),
        )
    _add_cell_borders(table)

    p = doc.add_paragraph()  # type: ignore[attr-defined]
    p.add_run(
        f"Total CGT payable across all years: ${tax.total_tax_paid:,.2f}   |   "
        "Note: after-tax CAGR not computed (requires price history)."
    ).font.size = Pt(9)
    doc.add_paragraph()  # type: ignore[attr-defined]


def _add_cgt_event_log_from_lots(doc: object, lots: list[DisposedLot], entity_type: str = "individual") -> None:
    """CGT event log from DisposedLot list (no TaxAwareResult needed)."""
    _section_heading(doc, "CGT Event Log")

    rows = build_cgt_event_rows(lots)
    if not rows:
        doc.add_paragraph("No disposal events recorded in this period.")  # type: ignore[attr-defined]
        return

    table = doc.add_table(rows=1, cols=7)  # type: ignore[attr-defined]
    _header_row(
        table,
        ["#", "Tax Year", "Ticker", "Acquired", "Disposed", "Gain/Loss (AUD)", "ATO Rule Applied"],
        font_size=8,
    )
    for i, row in enumerate(rows):
        _body_row(
            table,
            [
                str(i + 1),
                row.tax_year_label,
                row.ticker,
                row.acquired_date.strftime("%d %b %Y"),
                row.disposed_date.strftime("%d %b %Y"),
                f"${row.gain_aud:,.2f}",
                _cgt_rule_annotation(row, entity_type),
            ],
            shade=(i % 2 == 1),
            font_size=8,
        )
    _set_col_widths(table, [0.4, 1.5, 1.5, 2.0, 2.0, 2.1, 6.5])
    _add_cell_borders(table)
    doc.add_paragraph()  # type: ignore[attr-defined]


def export_trades_cgt_workpaper(
    trades: list[TradeRecord],
    tax: TaxSummary,
    output_path: Path,
    *,
    entity_type: str = "individual",
    broker: str = "unknown",
) -> str:
    """Export a CGT workpaper generated from actual broker trades (no backtest).

    Produces a Word document suitable for SMSF auditors. Includes the actual
    trade history, CGT event log, and year-by-year tax summary. Performance
    metrics (CAGR, Sharpe, drawdown) are omitted — they require price history
    which is not fetched in the actual-trades flow.

    Args:
        trades: Validated TradeRecord list from parse_broker_csv().
        tax: TaxSummary from run_cgt_from_trades().
        output_path: Destination .docx path. Parent must exist.
        entity_type: "individual" or "smsf" — shown on cover page.
        broker: Broker name for display (e.g. "commsec").

    Raises:
        ValueError: If output_path does not have a .docx extension.
    """
    if output_path.suffix.lower() != ".docx":
        raise ValueError(f"output_path must end in .docx, got: {output_path}")

    workpaper_id = generate_workpaper_id()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_page_numbers(doc)

    # Cover page
    title_para = doc.add_paragraph()  # type: ignore[attr-defined]
    title_run = title_para.add_run("PortfolioForge — CGT Workpaper (Actual Trades)")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _NAVY

    sorted_trades = sorted(trades, key=lambda t: t.trade_date)
    period_start = sorted_trades[0].trade_date if sorted_trades else date.today()
    period_end = sorted_trades[-1].trade_date if sorted_trades else date.today()
    tickers = sorted({r.ticker for r in trades})

    sub = doc.add_paragraph()  # type: ignore[attr-defined]
    sub.add_run(f"Securities: {', '.join(tickers)}").font.size = Pt(10)

    period = doc.add_paragraph()  # type: ignore[attr-defined]
    period.add_run(
        f"Trade period: {period_start.strftime('%d %b %Y')} to {period_end.strftime('%d %b %Y')}"
        f"   |   Entity: {entity_type.upper()}"
        f"   |   Broker: {broker.title()}"
        f"   |   Generated: {date.today().strftime('%d %b %Y')}"
    ).font.size = Pt(10)

    vid_para = doc.add_paragraph()  # type: ignore[attr-defined]
    vid_run = vid_para.add_run(f"Verification ID: {workpaper_id}")
    vid_run.font.size = Pt(9)
    vid_run.font.color.rgb = _NAVY

    doc.add_paragraph()  # type: ignore[attr-defined]

    # KPI row: 3 boxes
    kpi_table = doc.add_table(rows=2, cols=3)  # type: ignore[attr-defined]
    kpi_labels = ["Trades Imported", "CGT Events", "Total CGT Payable"]
    cgt_events = sum(yr.cgt_events for yr in tax.years)
    kpi_values = [str(len(trades)), str(cgt_events), f"${tax.total_tax_paid:,.2f}"]
    label_row = kpi_table.rows[0]
    value_row = kpi_table.rows[1]
    for i, (lbl, val) in enumerate(zip(kpi_labels, kpi_values, strict=True)):
        _shade_cell(label_row.cells[i], _NAVY_HEX)
        lp = label_row.cells[i].paragraphs[0]
        lp.clear()
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.color.rgb = _WHITE
        lr.font.size = Pt(9)
        _shade_cell(value_row.cells[i], _TEAL_HEX)
        vp = value_row.cells[i].paragraphs[0]
        vp.clear()
        vr = vp.add_run(val)
        vr.bold = True
        vr.font.color.rgb = _WHITE
        vr.font.size = Pt(14)

    doc.add_page_break()  # type: ignore[attr-defined]

    _add_trade_history_table(doc, trades)
    _add_cgt_tax_summary(doc, tax)
    _add_cgt_event_log_from_lots(doc, tax.lots, entity_type=entity_type)
    _add_methodology(doc, tax_result=None)
    _add_disclaimer(doc)

    doc.save(str(output_path))
    return workpaper_id
