"""Tests for analysis/exporter.py -- Word document export."""

from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from market_data.analysis.exporter import export_report
from market_data.verification.workpaper_id import verify_workpaper_id
from market_data.analysis.models import AnalysisReport
from market_data.analysis.narrative import DISCLAIMER
from market_data.backtest.models import (
    BacktestResult,
    BenchmarkResult,
    DataCoverage,
    PerformanceMetrics,
)
from market_data.backtest.tax.models import (
    DisposedLot,
    TaxAwareResult,
    TaxSummary,
    TaxYearResult,
)
from market_data.db.schema import get_connection

# ── fixtures ─────────────────────────────────────────────────────────────────


def _make_backtest(days: int = 60) -> BacktestResult:
    idx = pd.date_range("2022-01-01", periods=days, freq="D")
    equity = pd.Series([10_000.0 + i * 50 for i in range(days)], index=idx)
    bench = pd.Series([10_000.0 + i * 40 for i in range(days)], index=idx)
    return BacktestResult(
        metrics=PerformanceMetrics(
            total_return=0.30,
            cagr=0.12,
            max_drawdown=-0.08,
            sharpe_ratio=1.2,
        ),
        benchmark=BenchmarkResult(
            ticker="STW.AX",
            total_return=0.22,
            cagr=0.09,
            max_drawdown=-0.10,
            sharpe_ratio=0.9,
        ),
        equity_curve=equity,
        benchmark_curve=bench,
        trades=[],
        coverage=[
            DataCoverage(
                ticker="VAS.AX",
                from_date=date(2022, 1, 1),
                to_date=date(2022, 12, 31),
                records=days,
            )
        ],
        portfolio={"VAS.AX": 0.6, "VGS.AX": 0.4},
        initial_capital=10_000.0,
        start_date=date(2022, 1, 1),
        end_date=date(2022, 12, 31),
    )


def _make_tax_result(br: BacktestResult) -> TaxAwareResult:
    lot = DisposedLot(
        ticker="VAS.AX",
        acquired_date=date(2022, 1, 15),
        disposed_date=date(2022, 11, 20),
        quantity=100.0,
        cost_basis_usd=None,
        cost_basis_aud=Decimal("5000.0"),
        proceeds_usd=None,
        proceeds_aud=5_800.0,
        gain_aud=800.0,
        discount_applied=False,
    )
    yr = TaxYearResult(
        ending_year=2023,
        cgt_events=1,
        cgt_payable=360.0,
        franking_credits_claimed=120.0,
        dividend_income=400.0,
        after_tax_return=440.0,
        carried_forward_loss=0.0,
    )
    tax = TaxSummary(
        years=[yr],
        total_tax_paid=360.0,
        after_tax_cagr=0.09,
        lots=[lot],
        marginal_tax_rate=0.325,
    )
    return TaxAwareResult(backtest=br, tax=tax)


def _all_text(doc_path: Path) -> str:
    """Extract all text (paragraphs + table cells) from the document."""
    doc = Document(str(doc_path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return " ".join(parts)


# ── tests ─────────────────────────────────────────────────────────────────────


def test_export_creates_file(tmp_path: Path) -> None:
    """export_report() writes a .docx file to the given path."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "report.docx"

    export_report(report, conn, out)

    assert out.exists()
    assert out.stat().st_size > 0


def test_export_rejects_non_docx(tmp_path: Path) -> None:
    """export_report() raises ValueError for non-.docx extensions."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")

    with pytest.raises(ValueError, match=".docx"):
        export_report(report, conn, tmp_path / "report.pdf")


def test_disclaimer_always_present(tmp_path: Path) -> None:
    """The DISCLAIMER text appears in every exported document."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "report.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    assert DISCLAIMER in text


def test_tax_sections_included_for_tax_result(tmp_path: Path) -> None:
    """Tax summary and CGT log sections appear when TaxAwareResult is present."""
    br = _make_backtest()
    tax_result = _make_tax_result(br)
    report = AnalysisReport(result=tax_result)
    conn = get_connection(":memory:")
    out = tmp_path / "report.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    assert "Australian Tax Analysis" in text
    assert "CGT Event Log" in text
    # FY year from our fixture
    assert "FY2023" in text
    # CGT log uses 7-column format including ATO rule annotation (PROF-01)
    assert "Gain/Loss (AUD)" in text
    assert "ATO Rule Applied" in text
    assert "FIFO parcel" in text  # annotation text generated for each event
    # Calculation Methodology table present with ATO references (PROF-02)
    assert "Calculation Methodology" in text
    assert "ITAA 1997" in text
    assert "Marginal tax rate applied" in text


def test_tax_note_shown_for_pre_tax_result(tmp_path: Path) -> None:
    """A plain BacktestResult produces a note that tax data is unavailable."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "report.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    assert "Tax analysis not available" in text


def test_methodology_table_static_rules_always_present(tmp_path: Path) -> None:
    """Calculation Methodology ATO rules table appears even without tax data (PROF-02)."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "report.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    assert "Calculation Methodology" in text
    assert "FIFO" in text
    assert "ITAA 1997" in text
    # Dynamic tax rows should NOT appear for a pre-tax result
    assert "Marginal tax rate applied" not in text


def test_carry_forward_loss_in_tax_table(tmp_path: Path) -> None:
    """Carried-forward loss column appears in tax table when non-zero."""
    # Build a result with a non-zero carry-forward
    yr = TaxYearResult(
        ending_year=2023,
        cgt_events=1,
        cgt_payable=0.0,
        franking_credits_claimed=0.0,
        dividend_income=0.0,
        after_tax_return=-500.0,
        carried_forward_loss=500.0,
    )
    tax = TaxSummary(
        years=[yr],
        total_tax_paid=0.0,
        after_tax_cagr=-0.05,
        lots=[],
        marginal_tax_rate=0.325,
    )
    tax_result = TaxAwareResult(
        backtest=_make_backtest(),
        tax=tax,
    )
    report = AnalysisReport(result=tax_result)
    conn = get_connection(":memory:")
    out = tmp_path / "carry_forward.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    # Carry-forward column header present
    assert "Carry-Fwd Loss" in text
    # The value $500.00 appears
    assert "$500.00" in text


def test_sample_data_label_appears_when_flag_set(tmp_path: Path) -> None:
    """SAMPLE DATA banner appears on the cover when sample_data=True."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "sample.docx"

    export_report(report, conn, out, sample_data=True)

    text = _all_text(out)
    assert "SAMPLE DATA" in text


def test_sample_data_label_absent_by_default(tmp_path: Path) -> None:
    """No SAMPLE DATA banner when sample_data is not set (default False)."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "normal.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    assert "SAMPLE DATA" not in text


def test_coverage_quality_status_column_present(tmp_path: Path) -> None:
    """Data Coverage table includes a Quality Status column showing 'validated'."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "coverage.docx"

    export_report(report, conn, out)

    text = _all_text(out)
    assert "Quality Status" in text
    assert "validated" in text


# ── semantic tests (HARD-08) ──────────────────────────────────────────────────


def test_disclaimer_present_semantic(tmp_path: Path) -> None:
    """DISCLAIMER text appears in the exported document (structural check)."""
    br = _make_backtest()
    tax = _make_tax_result(br)
    out = tmp_path / "report.docx"
    conn = get_connection(":memory:")
    export_report(AnalysisReport(result=tax), conn, out)
    doc = Document(str(out))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert DISCLAIMER[:50] in all_text  # first 50 chars of disclaimer constant


def test_cgt_table_row_count_semantic(tmp_path: Path) -> None:
    """CGT summary table has exactly header + N data rows (one per tax year)."""
    br = _make_backtest()
    tax = _make_tax_result(br)
    out = tmp_path / "report.docx"
    conn = get_connection(":memory:")
    export_report(AnalysisReport(result=tax), conn, out)
    doc = Document(str(out))
    # The tax analysis table is the only 6-column table in the document.
    tax_tables = [t for t in doc.tables if len(t.columns) == 6]
    assert len(tax_tables) == 1, f"Expected 1 six-column table, got {len(tax_tables)}"
    # Fixture has 1 tax year -> header row + 1 data row = 2 total
    assert len(tax_tables[0].rows) == 2, (
        f"Expected 2 rows (header + 1 year), got {len(tax_tables[0].rows)}"
    )


def test_methodology_table_present_semantic(tmp_path: Path) -> None:
    """Methodology section creates a 3-column table in the document."""
    br = _make_backtest()
    out = tmp_path / "report.docx"
    conn = get_connection(":memory:")
    export_report(AnalysisReport(result=br), conn, out)
    doc = Document(str(out))
    three_col_tables = [t for t in doc.tables if len(t.columns) == 3]
    # Methodology table is the last 3-column table (Composition and Performance also have 3 cols)
    assert len(three_col_tables) >= 3, (
        f"Expected at least 3 three-column tables (Composition, Performance, Methodology), "
        f"got {len(three_col_tables)}"
    )


def test_methodology_table_row_count_semantic(tmp_path: Path) -> None:
    """Methodology table has at least header + 8 static rule rows."""
    br = _make_backtest()
    tax = _make_tax_result(br)
    out = tmp_path / "report.docx"
    conn = get_connection(":memory:")
    export_report(AnalysisReport(result=tax), conn, out)
    doc = Document(str(out))
    three_col_tables = [t for t in doc.tables if len(t.columns) == 3]
    assert len(three_col_tables) >= 3
    # Methodology table is the last 3-column table in the document
    methodology_table = three_col_tables[-1]
    # _CGT_RULES has 8 entries + header = 9 minimum
    # With tax_result (HARD-02 adds version row, plus 2 existing dynamic rows) = 12
    assert len(methodology_table.rows) >= 9, (
        f"Expected >= 9 rows in Methodology table, got {len(methodology_table.rows)}"
    )


# ── export_trades_cgt_workpaper() tests ───────────────────────────────────────


from market_data.analysis.exporter import export_trades_cgt_workpaper  # noqa: E402
from market_data.backtest.tax.engine import run_cgt_from_trades  # noqa: E402
from market_data.backtest.tax.trade_record import TradeRecord  # noqa: E402


def _make_tr(
    trade_date: date,
    ticker: str,
    action: str,
    quantity: float,
    price_aud: float,
    brokerage_aud: float = 0.0,
) -> TradeRecord:
    return TradeRecord(
        trade_date=trade_date,
        ticker=ticker,
        action=action,  # type: ignore[arg-type]
        quantity=quantity,
        price_aud=price_aud,
        brokerage_aud=brokerage_aud,
    )


def _simple_trades_and_tax() -> tuple[list[TradeRecord], object]:
    """One BUY + one short-term SELL: $750 gain, CGT $243.75."""
    trades = [
        _make_tr(date(2023, 1, 3), "VAS.AX", "BUY", 1000, 1.50, 50.0),
        _make_tr(date(2023, 6, 1), "VAS.AX", "SELL", 1000, 2.35, 50.0),
    ]
    tax = run_cgt_from_trades(trades, marginal_tax_rate=0.325)
    return trades, tax


def test_workpaper_creates_docx_file(tmp_path: Path) -> None:
    """export_trades_cgt_workpaper() writes a non-empty .docx file."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"

    export_trades_cgt_workpaper(trades, tax, out, entity_type="individual", broker="commsec")

    assert out.exists()
    assert out.stat().st_size > 0


def test_workpaper_rejects_non_docx_extension(tmp_path: Path) -> None:
    """Non-.docx extension raises ValueError."""
    trades, tax = _simple_trades_and_tax()

    with pytest.raises(ValueError, match=".docx"):
        export_trades_cgt_workpaper(trades, tax, tmp_path / "out.pdf")


def test_workpaper_is_parseable_docx(tmp_path: Path) -> None:
    """Resulting file can be opened by python-docx without error."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    # Should not raise
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


def test_workpaper_contains_portfolioforge_branding(tmp_path: Path) -> None:
    """Cover page includes 'PortfolioForge' branding."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "PortfolioForge" in text


def test_workpaper_disclaimer_always_present(tmp_path: Path) -> None:
    """DISCLAIMER constant must appear in every workpaper output."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert DISCLAIMER[:50] in text


def test_workpaper_entity_type_appears(tmp_path: Path) -> None:
    """Entity type is shown in uppercase on the cover page."""
    trades = [_make_tr(date(2023, 1, 3), "BHP.AX", "BUY", 50, 40.0)]
    tax = run_cgt_from_trades(trades, entity_type="smsf")
    out = tmp_path / "smsf.docx"

    export_trades_cgt_workpaper(trades, tax, out, entity_type="smsf", broker="selfwealth")

    text = _all_text(out)
    assert "SMSF" in text


def test_workpaper_broker_name_appears_title_cased(tmp_path: Path) -> None:
    """Broker name appears title-cased on the cover page."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"

    export_trades_cgt_workpaper(trades, tax, out, broker="commsec")

    text = _all_text(out)
    assert "Commsec" in text


def test_workpaper_ticker_appears_in_document(tmp_path: Path) -> None:
    """Security ticker appears in the trade history section."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "VAS.AX" in text


def test_workpaper_cgt_summary_section_present(tmp_path: Path) -> None:
    """'Australian Tax Analysis' section heading appears in workpaper."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "Australian Tax Analysis" in text


def test_workpaper_trade_history_section_present(tmp_path: Path) -> None:
    """Trade History section heading appears in workpaper."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "Trade History" in text


def test_workpaper_trade_history_table_row_count(tmp_path: Path) -> None:
    """Trade history table has exactly header + N data rows (one per trade)."""
    trades = [
        _make_tr(date(2023, 1, 3), "VAS.AX", "BUY", 100, 100.0, 50.0),
        _make_tr(date(2023, 3, 1), "VGS.AX", "BUY", 50, 200.0, 50.0),
        _make_tr(date(2023, 6, 1), "VAS.AX", "SELL", 100, 110.0, 50.0),
    ]
    tax = run_cgt_from_trades(trades)
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    doc = Document(str(out))
    # Trade history table has 7 columns.
    seven_col_tables = [t for t in doc.tables if len(t.columns) == 7]
    # At least one 7-column table (may share with CGT event log — also 7 cols).
    assert len(seven_col_tables) >= 1

    # The first 7-column table is the trade history: header + 3 data rows = 4 rows.
    trade_table = seven_col_tables[0]
    assert len(trade_table.rows) == 4, (
        f"Expected 4 rows (1 header + 3 trades), got {len(trade_table.rows)}"
    )


def test_workpaper_cgt_total_appears(tmp_path: Path) -> None:
    """Total CGT payable dollar amount appears in the workpaper."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    # CGT for this fixture = 750 * 0.325 = $243.75
    assert "243.75" in text


def test_workpaper_no_disposal_events_message(tmp_path: Path) -> None:
    """Buy-only trades produce 'No disposal events' message in CGT event log."""
    trades = [_make_tr(date(2023, 1, 3), "VAS.AX", "BUY", 100, 100.0)]
    tax = run_cgt_from_trades(trades)
    out = tmp_path / "no_disposals.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "No disposal events" in text


def test_workpaper_multiple_tickers_all_appear(tmp_path: Path) -> None:
    """All tickers in a multi-ticker trade list appear in the workpaper."""
    trades = [
        _make_tr(date(2023, 1, 3), "VAS.AX", "BUY", 100, 100.0),
        _make_tr(date(2023, 1, 3), "CBA.AX", "BUY", 50, 80.0),
        _make_tr(date(2023, 6, 1), "VAS.AX", "SELL", 100, 110.0),
        _make_tr(date(2023, 6, 1), "CBA.AX", "SELL", 50, 90.0),
    ]
    tax = run_cgt_from_trades(trades)
    out = tmp_path / "multi.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "VAS.AX" in text
    assert "CBA.AX" in text


def test_workpaper_methodology_section_present(tmp_path: Path) -> None:
    """Calculation Methodology section always appears."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "workpaper.docx"
    export_trades_cgt_workpaper(trades, tax, out)

    text = _all_text(out)
    assert "Calculation Methodology" in text
    assert "ITAA 1997" in text


# ── verification ID tests ─────────────────────────────────────────────────────


def test_export_report_returns_verification_id(tmp_path: Path) -> None:
    """export_report must return a valid workpaper ID string."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "out.docx"
    wid = export_report(report, conn, out)
    assert isinstance(wid, str)
    result = verify_workpaper_id(wid)
    assert result.valid is True


def test_export_trades_cgt_workpaper_returns_verification_id(tmp_path: Path) -> None:
    """export_trades_cgt_workpaper must return a valid workpaper ID string."""
    trades, tax = _simple_trades_and_tax()
    out = tmp_path / "out.docx"
    wid = export_trades_cgt_workpaper(trades, tax, out)
    assert isinstance(wid, str)
    result = verify_workpaper_id(wid)
    assert result.valid is True


def test_export_report_embeds_id_in_document(tmp_path: Path) -> None:
    """The verification ID must appear in the document text."""
    br = _make_backtest()
    report = AnalysisReport(result=br)
    conn = get_connection(":memory:")
    out = tmp_path / "out.docx"
    wid = export_report(report, conn, out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert wid in full_text
